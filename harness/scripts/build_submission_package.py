"""Build the MDPI submission package into submission_MDPI/.

    python -m harness.scripts.build_submission_package

Produces:
  submission_MDPI/graphical_abstract.png   1100 x 560 px (MDPI: "560 x 1100 (height x width)")
  submission_MDPI/figures.zip              contents of figures/, flat
  submission_MDPI/manuscript.zip           main.tex + bibliography.bib + main.bbl + Definitions/ + figures/
  submission_MDPI/reproduction.zip         harness + runs + figures + REPRODUCE.md  (= the Zenodo deposit)
  submission_MDPI/cover_letter.docx

Excluded from every archive, deliberately: .env (real API keys), __pycache__/*.pyc,
and harness/data/primevul/*.jsonl (the 66 MB PrimeVul corpus belongs to its original
authors and is re-downloaded rather than redistributed).
"""

from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import pandas as pd  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

ROOT = Path(".")
OUT = Path("submission_MDPI")
METRICS = Path("harness/runs/primevul_combined/enriched_metrics.csv")

TITLE = ("Open or Frontier? A Cost- and Energy-Aware Benchmark of "
         "Large Language Models for Software Vulnerability Detection")
JOURNAL = "Computers"
ZENODO_DOI = "10.5281/zenodo.21391074"

# Files whose presence in an archive would leak credentials or bloat the deposit.
BLOCKED_NAMES = {".env"}


def _skip(p: Path) -> bool:
    parts = set(p.parts)
    return (
        p.name in BLOCKED_NAMES
        or "__pycache__" in parts
        or ".pytest_cache" in parts
        or ".git" in parts
        or p.suffix in {".pyc", ".pyo"}
        # the PrimeVul corpus itself: redistributed by its authors, not by us
        or (p.match("harness/data/primevul/*"))
    )


def _add_tree(zf: zipfile.ZipFile, src: Path, arc_prefix: str = "") -> int:
    n = 0
    for p in sorted(src.rglob("*")):
        if p.is_dir() or _skip(p):
            continue
        arc = Path(arc_prefix) / p.relative_to(src) if arc_prefix else p
        zf.write(p, arc.as_posix())
        n += 1
    return n


# --------------------------------------------------------------------------- #
# Graphical abstract
# --------------------------------------------------------------------------- #
def graphical_abstract(df: pd.DataFrame, out: Path) -> None:
    """1100 x 560 px. MDPI states the minimum as '560 x 1100 pixels (height x width)',
    i.e. landscape: 1100 wide, 560 tall."""
    px_w, px_h, dpi = 1100, 560, 100
    fig = plt.figure(figsize=(px_w / dpi, px_h / dpi), dpi=dpi)
    fig.patch.set_facecolor("white")
    blue, red, ink, mute = "#1f77b4", "#d62728", "#1a1a1a", "#5c5c5c"

    fig.text(0.035, 0.915, "Open-weight LLMs own the efficiency frontier",
             fontsize=19, fontweight="bold", color=ink, va="top")
    fig.text(0.035, 0.845,
             "Eight LLMs on function-level vulnerability detection  ·  PrimeVul, N = 1549  ·  "
             "cost and energy as first-class axes",
             fontsize=10.5, color=mute, va="top")

    # ---- left: the energy Pareto panel -------------------------------------
    # Left inset leaves room for the y-label; the right pad is deliberately wider so
    # Claude-Sonnet-5 (the rightmost point) can be labelled outward instead of having
    # its label run back across the Llama-3.3-70B marker.
    ax = fig.add_axes([0.085, 0.16, 0.37, 0.58])
    for tier, c in (("open", blue), ("frontier", red)):
        s = df[df["tier"] == tier]
        ax.scatter(s["energy_j"], s["bal_acc"], c=c, s=70, zorder=3, label=tier)
    ax.set_xscale("log")
    lo, hi = df["energy_j"].min(), df["energy_j"].max()
    ax.set_xlim(lo / (hi / lo) ** 0.12, hi * (hi / lo) ** 0.52)
    ylo, yhi = ax.get_ylim()
    ax.set_ylim(ylo, yhi + 0.10 * (yhi - ylo))
    for name, dy in (("Gemma-3-4B", 4), ("DeepSeek-V3.2", 4), ("Claude-Sonnet-5", 4)):
        r = df[df["model"] == name].iloc[0]
        ax.annotate(name, (r["energy_j"], r["bal_acc"]), textcoords="offset points",
                    xytext=(8, dy), fontsize=9, ha="left", va="bottom")
    ax.axhline(0.5, ls="--", c="grey", lw=1, zorder=1)
    ax.set_xlabel("Energy per task, J (log)", fontsize=10)
    ax.set_ylabel("Balanced accuracy", fontsize=10)
    ax.tick_params(labelsize=9)
    ax.legend(frameon=False, fontsize=9.5, loc="lower right")

    # ---- right: the three headline numbers ---------------------------------
    g = df[df["model"] == "Gemma-3-4B"].iloc[0]
    c = df[df["model"] == "Claude-Sonnet-5"].iloc[0]
    cost_x = c["usd_task"] / g["usd_task"]
    energy_x = c["energy_j"] / g["energy_j"]

    # No "$" anywhere in these strings: matplotlib reads it as mathtext and mangles the line.
    rows = [
        (f"{cost_x:.0f}×", "lower cost per task",
         f"Gemma-3-4B USD {g['usd_task']:.5f} vs Claude-Sonnet-5 USD {c['usd_task']:.4f}"),
        (f"{energy_x:.0f}×", "less energy per task",
         f"{g['energy_j']:.0f} J vs {c['energy_j']:.0f} J  (10–150× across assumptions)"),
        ("Top 3", "on quality are all open-weight",
         "and no frontier model is Pareto-optimal"),
    ]
    x_num, x_txt, y = 0.525, 0.645, 0.665
    for big, lead, sub in rows:
        fig.text(x_num, y, big, fontsize=25, fontweight="bold", color=blue,
                 va="center", ha="left")
        fig.text(x_txt, y + 0.045, lead, fontsize=11.5, color=ink, va="center", ha="left")
        fig.text(x_txt, y - 0.048, sub, fontsize=9, color=mute, va="center", ha="left")
        y -= 0.215

    fig.text(0.525, 0.075,
             "A 4B open model matches or beats frontier quality at ~1/100 the cost\n"
             "and one to two orders of magnitude less energy per task.",
             fontsize=10, color=ink, va="bottom", ha="left", linespacing=1.5)

    fig.savefig(out, dpi=dpi, facecolor="white")
    plt.close(fig)
    from PIL import Image
    w, h = Image.open(out).size
    assert (w, h) == (px_w, px_h), f"graphical abstract is {w}x{h}, expected {px_w}x{px_h}"


# --------------------------------------------------------------------------- #
# Cover letter
# --------------------------------------------------------------------------- #
def cover_letter(out: Path) -> None:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(10)

    def para(text, bold=False, align=None, size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if align:
            p.alignment = align
        return p

    para(date.today().strftime("%d %B %Y"), align=WD_ALIGN_PARAGRAPH.RIGHT)
    para("The Editorial Office\nComputers (ISSN 2073-431X)\nMDPI, Basel, Switzerland")
    para("Dear Editors,")

    para(
        f"Please find enclosed our manuscript entitled “{TITLE}”, which we submit "
        f"for consideration as a research article in {JOURNAL}. The work is original, has not "
        "been published previously, and is not under consideration for publication elsewhere.")

    para("Motivation and gap", bold=True)
    para(
        "Large language models are now routinely evaluated on software vulnerability detection, "
        "but that literature has two blind spots. First, it is almost exclusively accuracy-centric: "
        "an organisation running triage over a large codebase also needs to know what each verdict "
        "costs in money and in energy, yet the environmental and economic footprint of inference is "
        "essentially unreported in security benchmarking. Second, it leans heavily on proprietary "
        "frontier models, even though open-weight models are precisely what many security teams need "
        "in order to keep source code on-premises for privacy, compliance and data-residency reasons. "
        "Our manuscript brings these two concerns together.")

    para("What we did", bold=True)
    para(
        "We benchmark eight contemporary LLMs — three frontier and five open-weight — on a "
        "stratified 1549-function sample of the label-clean PrimeVul dataset, treating cost per task "
        "and energy as first-class evaluation axes alongside detection quality. Monetary cost and "
        "latency are measured directly for every model. Energy is measured directly on a dedicated "
        "NVIDIA H200 GPU via the NVML counter for the two models we could serve locally, and estimated "
        "with a transparent, assumption-explicit FLOP model for the API-served models; every result row "
        "carries a tag so measured and estimated energy are never conflated. We calibrate the LLMs "
        "against a classical static analyser, an off-the-shelf learned detector, and both trivial "
        "classifiers.")

    para("Principal findings", bold=True)
    para(
        "Open-weight models occupy the quality and efficiency Pareto frontier, and no frontier model is "
        "Pareto-optimal. A 4-billion-parameter open model matches or exceeds frontier detection quality "
        "at roughly one-hundredth the monetary cost and one to two orders of magnitude less energy per "
        "task. The three highest-scoring models on balanced accuracy are all open-weight, and on a "
        "Holm-corrected paired bootstrap the best of them significantly exceeds all three frontier "
        "models in the direct-answer configuration. Allowing the frontier models their native reasoning "
        "mode closes that quality gap to statistical parity but costs 2.5 to 7.6 times more, which "
        "widens rather than narrows the efficiency separation. We are equally explicit about what the "
        "benchmark does not show: absolute quality is modest throughout, and at PrimeVul’s realistic "
        "1:44 prevalence every model yields low precision, so the finding is a conditional deployment "
        "recommendation rather than a claim that the task is solved.")

    para(f"Fit with {JOURNAL}", bold=True)
    para(
        f"{JOURNAL} publishes work at the intersection of applied computing, software engineering and "
        "security, and has a standing interest in the practical and sustainable deployment of machine "
        "learning systems. Our contribution is a measurement study that imports the rigour of "
        "energy-aware benchmarking into a security-detection task, which to our knowledge has not been "
        "done with this combination of open-weight breadth and per-task cost and energy accounting. We "
        "believe it will be of direct use to readers who must choose and deploy such tooling.")

    para("Reproducibility and data availability", bold=True)
    para(
        "All experiments run through a purpose-built, test-driven measurement harness (82 automated "
        "tests). The harness, all run configurations, the exact sampled function identifiers, and the "
        "aggregated results are openly deposited at Zenodo "
        f"(https://doi.org/{ZENODO_DOI}). The PrimeVul dataset is available from its original authors.")

    para("Declarations", bold=True)
    para(
        "In accordance with journal policy, the manuscript carries a “Use of Artificial Intelligence” "
        "statement in the back matter. In brief: DeepL and ChatGPT were used to translate words, "
        "phrases and passages from German into English, and Claude was used to assist with drafting and "
        "revising portions of the text and with writing Python code. No text was published without "
        "author review, and no figures, data or images were AI-generated or fabricated — all figures "
        "derive from our own measurements, produced by the released harness from logged data. The "
        "authors have reviewed and edited all outputs and take full responsibility for the content.")
    para(
        "All authors have read and agreed to the submitted version of the manuscript. The authors "
        "declare no conflicts of interest. This research was funded in part by the State of Styria (Land "
        "Steiermark), Office of the Styrian Provincial Government, Department 12, within the PRISMA "
        "project, grant number ABT12-270413/2024. The funders had no role in the design of the study, "
        "in the collection, analyses or interpretation of data, in the writing of the manuscript, or in "
        "the decision to publish the results.")

    para("We look forward to hearing from you and would be glad to respond to any questions.")
    para("Yours sincerely,")
    para(
        "Patrick Deininger (corresponding author)\n"
        "Institute of Software Engineering and Artificial Intelligence,\n"
        "Graz University of Technology, 8010 Graz, Austria\n"
        "Institute of Computer Science and Artificial Intelligence,\n"
        "FH JOANNEUM – University of Applied Sciences, 8020 Graz, Austria\n"
        "patrick.deininger@student.tugraz.at\n\n"
        "on behalf of: Patrick Deininger and Wolfgang Slany")

    doc.save(out)


# --------------------------------------------------------------------------- #
REPRODUCE_MD = f"""# Reproduction package

Measurement harness and results for the paper:

**{TITLE}**
Patrick Deininger and Wolfgang Slany. Submitted to MDPI *{JOURNAL}*.

Archived at Zenodo: https://doi.org/{ZENODO_DOI}

## Contents

| Path | What it is |
|---|---|
| `harness/` | The measurement harness: backends, energy meters, scoring, analysis, 82 tests |
| `harness/runs/` | Every run's raw per-task records, resolved config, seed and price snapshot |
| `harness/runs/primevul_combined/` | The main N=1549 run behind Table 2 and both figures |
| `harness/configs/` | Run configurations |
| `harness/scripts/` | Analysis entry points (see below) |
| `figures/` | The two Pareto figures as they appear in the paper |
| `main.pdf` | The manuscript |
| `LICENSE` | CC BY 4.0, plus the carve-outs for third-party material |

**Not included:** `harness/data/primevul/primevul_test.jsonl` (66 MB). The PrimeVul corpus
belongs to its original authors and is re-downloaded rather than redistributed — see
`harness/README.md`. Everything needed to re-run the *analysis* is present; only re-running
inference against the models requires the corpus.

## Setup

```bash
python -m venv .venv && . .venv/bin/activate     # Windows: .venv\\Scripts\\activate
pip install -r harness/requirements.txt
python -m pytest harness/tests -q                # expect 82 passed
```

## Reproduce the analysis (no API keys, no GPU, no spend)

Every number in the paper is derived from the committed run records:

```bash
python -m harness.scripts.build_final_analysis   # Table 2, both Pareto figures, enriched_metrics.csv
python -m harness.scripts.revision_stats         # bootstrap CIs + Holm-corrected paired tests
python -m harness.scripts.reasoning_vs_direct    # Section 4.5
python -m harness.scripts.contamination_check    # CVE-year memorisation gradient
python -m harness.scripts.flawfinder_baseline    # static-analysis reference point
python -m harness.scripts.learned_baseline       # CodeBERT-Devign reference point
```

Bootstrap results are seeded (seed 12345, n_boot 20000) and reproduce exactly.

## Re-run inference (costs money)

Requires the PrimeVul test split placed at `harness/data/primevul/primevul_test.jsonl` and
API credentials. Copy `.env.example` to `.env` and fill in your own keys — no credentials
are included in this archive.

```bash
python -m harness.run --config harness/configs/<config>.yaml
```

Generation is deterministic (temperature 0, fixed maximum output length) and results are
written incrementally, so long runs are crash-safe and resumable.

## Measured energy

On-GPU energy for the two locally servable open models was measured on a dedicated NVIDIA
H200 via the NVML cumulative energy counter at concurrency 1, gross and idle-subtracted.
See `harness/README.md` for the cloud-GPU procedure.
"""


def main() -> int:
    OUT.mkdir(exist_ok=True)
    df = pd.read_csv(METRICS)

    ga = OUT / "graphical_abstract.png"
    graphical_abstract(df, ga)
    print(f"  {ga}  (1100x560)")

    with zipfile.ZipFile(OUT / "figures.zip", "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(Path("figures").iterdir()):
            if p.is_file() and not _skip(p):
                zf.write(p, p.name)
        n_fig = len(zf.namelist())
    print(f"  {OUT / 'figures.zip'}  ({n_fig} files)")

    with zipfile.ZipFile(OUT / "manuscript.zip", "w", zipfile.ZIP_DEFLATED) as zf:
        for f in ("main.tex", "bibliography.bib", "main.bbl", "main.pdf"):
            zf.write(f, f)
        _add_tree(zf, Path("Definitions"), "Definitions")
        _add_tree(zf, Path("figures"), "figures")
        n_ms = len(zf.namelist())
    print(f"  {OUT / 'manuscript.zip'}  ({n_ms} files)")

    with zipfile.ZipFile(OUT / "reproduction.zip", "w", zipfile.ZIP_DEFLATED) as zf:
        _add_tree(zf, Path("harness"), "harness")
        _add_tree(zf, Path("figures"), "figures")
        zf.write("pyproject.toml", "pyproject.toml")
        zf.write(".env.example", ".env.example")
        zf.write("main.pdf", "main.pdf")
        zf.write("LICENSE", "LICENSE")
        zf.writestr("REPRODUCE.md", REPRODUCE_MD)
        n_rep = len(zf.namelist())
    print(f"  {OUT / 'reproduction.zip'}  ({n_rep} files)")

    # Keep the loose copy in the submission folder identical to the canonical root file.
    (OUT / "LICENSE").write_bytes(Path("LICENSE").read_bytes())
    print(f"  {OUT / 'LICENSE'}")

    cl = OUT / "cover_letter.docx"
    cover_letter(cl)
    print(f"  {cl}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
